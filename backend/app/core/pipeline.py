import os
import re
import shutil
import subprocess
import tempfile
import glob
import base64
import io
import textwrap
import time
from pathlib import Path
from typing import Optional
from functools import wraps

from PIL import Image
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4, landscape as rl_landscape
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas


def retry_with_backoff(max_retries=3, base_delay=10):
    """Decorator: retry on 429/503 with exponential backoff."""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            for attempt in range(max_retries + 1):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    msg = str(e)
                    if "503" in msg or "429" in msg or "rate" in msg.lower() or "unavailable" in msg.lower():
                        if attempt == max_retries:
                            raise
                        delay = base_delay * (2 ** attempt)
                        time.sleep(delay)
                    else:
                        raise
        return wrapper
    return decorator


DOCX_TEMPLATE_HEADER = """**GRACE HOUSE INTERNATIONAL SCHOOL**
**THIRD TERMINAL EXAMINATION**
**(2025/2026 SESSION)**
"""

OCR_PROMPT_TEMPLATE = """You are an expert academic typist and OCR specialist. Transcribe the handwritten examination sheets from the provided images.

OUTPUT FORMAT — Follow this structure exactly, no deviations:

1. Start with these exact three header lines:
**GRACE HOUSE INTERNATIONAL SCHOOL**
**THIRD TERMINAL EXAMINATION**
**(2025/2026 SESSION)**

2. Metadata line (NOT bold, no markdown):
SUBJECT: {subject}      CLASS: {class_name}

3. Instructions line (bold):
**INSTRUCTIONS: {instructions}**

4. Leave one blank line, then section headers in ALL CAPS BOLD:
**SECTION A (MULTIPLE CHOICE)**
Leave one blank line before and after section headers.

5. Multiple choice: Put options on the SAME LINE as the question, separated by single spaces.
Use lowercase letters in parentheses: (a) (b) (c) (d)
Example: 1. Which game uses a racket and ball? (a) Football (b) Table Tennis (c) Swimming

6. Theory questions: Use sub-letters indented under the question number.
Example:
1 a) What is noise pollution?
   b) List 5 causes of noise pollution.

7. Fix obvious spelling errors based on context (e.g., "fether" → "father", "freg" → "frog").

8. Return ONLY the formatted text. No explanations, no greetings, no extra commentary."""


@retry_with_backoff(max_retries=3, base_delay=10)
def ocr_images_gemini(images: list[str], api_key: str, subject: str, class_name: str,
                      instructions: str = "", model_name: str = "gemini-2.5-flash") -> str:
    """Send images to Gemini, return formatted OCR text."""
    from google import genai
    prompt = OCR_PROMPT_TEMPLATE.format(
        subject=subject, class_name=class_name,
        instructions=instructions or "Answer Section A and any TWO questions from Section B.",
    )
    client = genai.Client(api_key=api_key)
    contents = [prompt]
    for img_path in images:
        img = Image.open(img_path)
        contents.append(img)
    response = client.models.generate_content(model=model_name, contents=contents)
    return response.text


@retry_with_backoff(max_retries=3, base_delay=10)
def ocr_images_openrouter(images: list[str], api_key: str, subject: str, class_name: str,
                          instructions: str = "", model: str = "openai/gpt-4o-mini") -> str:
    """Send images to OpenRouter, return formatted OCR text."""
    import httpx
    import base64
    prompt = OCR_PROMPT_TEMPLATE.format(
        subject=subject, class_name=class_name,
        instructions=instructions or "Answer Section A and any TWO questions from Section B.",
    )
    content: list[dict] = [{"type": "text", "text": prompt}]
    for img_path in images:
        with open(img_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        ext = os.path.splitext(img_path)[1].lower()
        mime = "image/png" if ext == ".png" else "image/jpeg"
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{mime};base64,{b64}"},
        })
    body = {
        "model": model,
        "messages": [{"role": "user", "content": content}],
        "max_tokens": 8192,
    }
    with httpx.Client(timeout=120.0) as client:
        r = client.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://github.com/exam-platform",
                "X-Title": "Scribe Exam Platform",
            },
            json=body,
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]


def ocr_images(images: list[str], api_key: str, subject: str, class_name: str,
               instructions: str = "", model_name: str = "gemini-2.5-flash") -> str:
    """Dispatch to Gemini or OpenRouter based on API key pattern."""
    from app.config import settings
    if api_key == settings.openrouter_api_key or api_key.startswith("sk-or-"):
        return ocr_images_openrouter(images, api_key, subject, class_name, instructions, model_name)
    return ocr_images_gemini(images, api_key, subject, class_name, instructions, model_name)


def ocr_text_to_docx(text: str, subject: str, class_name: str, output_path: str) -> str:
    """Parse OCR text (with **bold** markers) and build a formatted DOCX.
    Mirrors the OCR text structure faithfully — headers get centered/bold treatment,
    everything else follows the original line structure with no extra spacing."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    lines = text.strip().split('\n')
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        is_bold = re.match(r'\*\*(.+?)\*\*', stripped)
        is_section_header = re.match(r'\*\*SECTION\s', stripped, re.I) or \
                            re.match(r'SECTION\s', stripped, re.I)

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        if i < 3 and not is_section_header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            txt = is_bold.group(1) if is_bold else stripped
            run = p.add_run(txt)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
        elif is_section_header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clean = stripped.replace('**', '')
            run = p.add_run(clean.upper())
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
        elif stripped.startswith('**INSTRUCTIONS'):
            run = p.add_run(stripped.replace('**', ''))
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
        elif stripped.startswith('SUBJECT:'):
            run = p.add_run(stripped)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
        else:
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                bm = re.match(r'\*\*(.*?)\*\*', part)
                if bm:
                    run = p.add_run(bm.group(1))
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

        i += 1

    doc.save(output_path)
    return output_path


def clone_run_format(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.strike = src_run.font.strike
    dst_run.font.superscript = src_run.font.superscript
    dst_run.font.subscript = src_run.font.subscript
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.color and src_run.font.color.rgb:
        try:
            dst_run.font.color.rgb = src_run.font.color.rgb
        except Exception:
            pass


def build_docx_from_ocr(
    text: str,
    subject: str,
    class_name: str,
    output_path: str,
    font_name: str = "Times New Roman",
    font_size_pt: int = 10,
) -> str:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    lines = text.strip().split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        is_bold = re.match(r"\*\*(.+?)\*\*", stripped)
        is_section_header = re.match(r"\*\*SECTION\s", stripped, re.I) or re.match(
            r"SECTION\s", stripped, re.I
        )

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        if i < 3 and not is_section_header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            txt = is_bold.group(1) if is_bold else stripped
            run = p.add_run(txt)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
        elif is_section_header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clean = stripped.replace("**", "")
            run = p.add_run(clean.upper())
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
        elif stripped.startswith("**INSTRUCTIONS"):
            run = p.add_run(stripped.replace("**", ""))
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
        elif stripped.startswith("SUBJECT:"):
            run = p.add_run(stripped)
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
        else:
            parts = re.split(r"(\*\*.*?\*\*)", stripped)
            for part in parts:
                bm = re.match(r"\*\*(.*?)\*\*", part)
                if bm:
                    run = p.add_run(bm.group(1))
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = font_name
                run.font.size = Pt(font_size_pt)

        i += 1

    doc.save(output_path)
    return output_path


def find_section_boundary(paragraphs):
    for i, p in enumerate(paragraphs):
        if re.match(r"(?i)^section\s+b", p.text.strip()):
            return i
    return None


def build_compact_docx(paragraphs, ranges, page_w_cm, page_h_cm, margin_cm, output_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(page_w_cm)
    sec.page_height = Cm(page_h_cm)
    for m in ["left_margin", "right_margin", "top_margin", "bottom_margin"]:
        setattr(sec, m, Cm(margin_cm))
    doc.styles["Normal"].font.size = Pt(12)

    for rng_start, rng_end in ranges:
        for pi in range(rng_start, rng_end):
            src = paragraphs[pi]
            text = src.text
            if not text.strip() and not src.runs:
                doc.add_paragraph("")
                continue
            is_hdr = pi < 3
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_hdr else src.alignment
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for run in src.runs:
                r = p.add_run(run.text)
                clone_run_format(run, r)
                if is_hdr:
                    r.bold = True
                if r.font.size is None:
                    r.font.size = Pt(12)
            if not src.runs and text.strip():
                r = p.add_run(text)
                if is_hdr:
                    r.bold = True
                r.font.size = Pt(12)

    doc.save(output_path)
    return output_path


def modify_docx_font_sizes(docx_path, scale_factor, output_path=None):
    if output_path is None:
        output_path = docx_path
    doc = Document(docx_path)
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size:
                r.font.size = Pt(int(r.font.size.pt * scale_factor))
    for sname in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]:
        if sname in doc.styles:
            s = doc.styles[sname]
            if s.font and s.font.size:
                s.font.size = Pt(int(s.font.size.pt * scale_factor))
    doc.save(output_path)
    return output_path


def docx_to_pdf(docx_path, output_pdf, timeout=120):
    outdir = os.path.dirname(output_pdf) or "."
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
        capture_output=True,
        text=True,
        timeout=timeout,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice failed: {result.stderr}")
    base = Path(docx_path).stem
    expected = os.path.join(outdir, f"{base}.pdf")
    if os.path.exists(expected):
        if expected != output_pdf:
            shutil.move(expected, output_pdf)
        return output_pdf
    candidates = glob.glob(os.path.join(outdir, "*.pdf"))
    if candidates:
        latest = max(candidates, key=os.path.getmtime)
        if latest != output_pdf:
            shutil.move(latest, output_pdf)
        return output_pdf
    raise RuntimeError("LibreOffice produced no PDF output")


def estimate_font_scale_and_render(
    docx_path: str,
    paragraphs,
    para_ranges: list[tuple[int, int]],
    page_w_cm: float,
    page_h_cm: float,
    margin_cm: float,
    target_pages: int = 1,
    max_target_pages: int | None = None,
    min_scale: float = 0.40,
) -> tuple[float, str]:
    """
    Render section DOCX once at scale=1.0, count pages, compute needed scale,
    apply it, render once more only if needed. Returns (scale_used, pdf_path).

    If max_target_pages is set and scale would be below 0.5, targets the larger
    page count instead — keeps font readable when content is only slightly long.

    Total LibreOffice calls: 1 if content already fits, 2 if scaling needed.
    Caller owns the returned pdf_path and must delete it when done.
    """
    tmp_docx_v1 = tempfile.mktemp(suffix=".docx")
    build_compact_docx(paragraphs, para_ranges, page_w_cm, page_h_cm, margin_cm, tmp_docx_v1)

    tmp_pdf_v1 = tempfile.mktemp(suffix=".pdf")
    docx_to_pdf(tmp_docx_v1, tmp_pdf_v1)
    os.unlink(tmp_docx_v1)

    pages = pdf_page_count_fitz(tmp_pdf_v1)

    if pages <= target_pages:
        return 1.0, tmp_pdf_v1

    scale = (target_pages / pages) ** (2.0 / 3.0)

    if max_target_pages and scale < 0.5 and max_target_pages > target_pages:
        scale = (max_target_pages / pages) ** (2.0 / 3.0)

    scale = max(min_scale, min(1.0, scale))

    os.unlink(tmp_pdf_v1)

    tmp_docx_v2 = tempfile.mktemp(suffix=".docx")
    build_compact_docx(paragraphs, para_ranges, page_w_cm, page_h_cm, margin_cm, tmp_docx_v2)
    modify_docx_font_sizes(tmp_docx_v2, scale, tmp_docx_v2)

    tmp_pdf_v2 = tempfile.mktemp(suffix=".pdf")
    docx_to_pdf(tmp_docx_v2, tmp_pdf_v2)
    os.unlink(tmp_docx_v2)

    actual_pages = pdf_page_count_fitz(tmp_pdf_v2)
    if actual_pages > target_pages:
        scale = max(min_scale, scale * 0.90)
        if actual_pages > target_pages + 1:
            os.unlink(tmp_pdf_v2)
            tmp_docx_v3 = tempfile.mktemp(suffix=".docx")
            build_compact_docx(paragraphs, para_ranges, page_w_cm, page_h_cm, margin_cm, tmp_docx_v3)
            modify_docx_font_sizes(tmp_docx_v3, scale, tmp_docx_v3)
            tmp_pdf_v3 = tempfile.mktemp(suffix=".pdf")
            docx_to_pdf(tmp_docx_v3, tmp_pdf_v3)
            os.unlink(tmp_docx_v3)
            return scale, tmp_pdf_v3

    return scale, tmp_pdf_v2


def pdf_page_to_image(pdf_path: str, page_num: int, dpi: int = 300) -> Image.Image:
    """Render a single PDF page to a PIL Image using PyMuPDF (no subprocess)."""
    import fitz
    doc = fitz.open(pdf_path)
    page = doc.load_page(page_num - 1)
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    doc.close()
    return img


def pdf_page_count_fitz(pdf_path: str) -> int:
    """Count PDF pages using PyMuPDF (no subprocess)."""
    import fitz
    doc = fitz.open(pdf_path)
    n = doc.page_count
    doc.close()
    return n


def impose_grid_pages(
    input_pdfs,
    output_pdf,
    cols=3,
    rows=2,
    margin_mm=2.5,
    gap_mm=2,
    cut_marks=True,
    fill_mode=True,
    labels=None,
):
    page_w, page_h = rl_landscape(A4)
    margin = margin_mm * mm
    gap = gap_mm * mm
    cell_w = (page_w - 2 * margin - (cols - 1) * gap) / cols
    cell_h = (page_h - 2 * margin - (rows - 1) * gap) / rows
    c = canvas.Canvas(output_pdf, pagesize=(page_w, page_h))
    cut_marks_drawn = False
    output_page_num = 0

    for src_idx, pdf_path in enumerate(input_pdfs):
        src_pages = pdf_page_count_fitz(pdf_path)

        for src_pg in range(1, src_pages + 1):
            if output_page_num > 0:
                c.showPage()
            output_page_num += 1

            img = pdf_page_to_image(pdf_path, src_pg, 200)
            tmp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            img.save(tmp_img, format="PNG")
            tmp_img.close()
            scale_val = max(cell_w / img.size[0], cell_h / img.size[1]) if fill_mode else min(
                cell_w / img.size[0], cell_h / img.size[1]
            )
            disp_w = img.size[0] * scale_val
            disp_h = img.size[1] * scale_val

            if cut_marks and not cut_marks_drawn:
                c.setStrokeColorRGB(0.82, 0.82, 0.82)
                c.setLineWidth(0.4)
                for gc in range(1, cols):
                    gx = margin + gc * cell_w + (gc - 1) * gap + gap / 2
                    c.line(gx, margin - 3, gx, page_h - margin + 3)
                for gr in range(1, rows):
                    gy = margin + gr * cell_h + (gr - 1) * gap + gap / 2
                    c.line(margin - 3, gy, page_w - margin + 3, gy)
                cut_marks_drawn = True

            for idx in range(cols * rows):
                row = idx // cols
                col = idx % cols
                py = rows - 1 - row
                cx = margin + col * (cell_w + gap)
                cy = margin + py * (cell_h + gap)
                x = cx + (cell_w - disp_w) / 2
                y = cy + (cell_h - disp_h) / 2

                if fill_mode and (disp_w > cell_w or disp_h > cell_h):
                    c.saveState()
                    p = c.beginPath()
                    p.rect(cx, cy, cell_w, cell_h)
                    c.clipPath(p, stroke=0)
                    c.drawImage(tmp_img.name, x, y, width=disp_w, height=disp_h)
                    c.restoreState()
                else:
                    c.drawImage(
                        tmp_img.name, x, y, width=disp_w, height=disp_h, preserveAspectRatio=True
                    )

                if labels and idx == 0:
                    label = labels[src_idx] if src_idx < len(labels) else ""
                    if src_pages > 1:
                        label += f" p{src_pg}"
                    c.setFont("Helvetica", 6)
                    c.setFillColorRGB(0.5, 0.5, 0.5)
                    c.drawString(cx + 2, cy + cell_h - 10, label)

            os.unlink(tmp_img.name)

    c.save()
    return output_pdf


def process_exam(
    docx_path: str,
    cols: int = 3,
    rows: int = 2,
    margin_mm: float = 2.5,
    gap_mm: float = 2.0,
    page_margin_cm: float = 0.25,
    split_mode: str = "Auto",
    header_pg2: bool = False,
    manual_scale_a: float = 0,
    manual_scale_b: float = 0,
    scale_a: int = 100,
    scale_b: int = 100,
) -> str:
    """
    DOCX → section PDFs → imposed grid PDF.

    estimate_font_scale_and_render() returns the PDF directly,
    so we never call docx_to_pdf() more than twice total (once per section).

    manual_scale_a / manual_scale_b: if > 0, use as direct font scale factor.
    scale_a / scale_b: percentage applied on top of auto scale (100 = no change).
    """
    doc = Document(docx_path)
    paragraphs = doc.paragraphs

    sec_b = None
    sec_a = None
    if split_mode == "Auto":
        for i, p in enumerate(paragraphs):
            t = p.text.strip()
            if re.match(r"(?i)^section\s+a", t) and sec_a is None:
                sec_a = i
            elif re.match(r"(?i)^section\s+b", t) and sec_b is None:
                sec_b = i

    pw, ph = rl_landscape(A4)
    mg = margin_mm * mm
    gp = gap_mm * mm
    cw = (pw - 2 * mg - (cols - 1) * gp) / cols
    ch = (ph - 2 * mg - (rows - 1) * gp) / rows
    cw_cm = cw / mm * 0.1
    ch_cm = ch / mm * 0.1

    pdf_a = None
    pdf_b = None

    try:
        if sec_b is not None:
            rng_a = [(0, sec_a), (sec_a, sec_b)]
            rng_b = [(sec_b, len(paragraphs))] if not header_pg2 else [(0, sec_a), (sec_b, len(paragraphs))]

            if manual_scale_a > 0:
                tmp_d = tempfile.mktemp(suffix=".docx")
                build_compact_docx(paragraphs, rng_a, cw_cm, ch_cm, page_margin_cm, tmp_d)
                modify_docx_font_sizes(tmp_d, manual_scale_a * (scale_a / 100), tmp_d)
                pdf_a = tempfile.mktemp(suffix=".pdf")
                docx_to_pdf(tmp_d, pdf_a)
                os.unlink(tmp_d)
            else:
                auto_sa, pdf_a = estimate_font_scale_and_render(
                    docx_path, paragraphs, rng_a,
                    cw_cm, ch_cm, page_margin_cm,
                    target_pages=1, max_target_pages=2,
                )
                if scale_a != 100:
                    adj = auto_sa * (scale_a / 100.0)
                    adj = max(0.35, min(5.0, adj))
                    if abs(adj - auto_sa) > 0.03:
                        os.unlink(pdf_a)
                        tmp_d = tempfile.mktemp(suffix=".docx")
                        build_compact_docx(paragraphs, rng_a, cw_cm, ch_cm, page_margin_cm, tmp_d)
                        modify_docx_font_sizes(tmp_d, adj, tmp_d)
                        pdf_a = tempfile.mktemp(suffix=".pdf")
                        docx_to_pdf(tmp_d, pdf_a)
                        os.unlink(tmp_d)

            if manual_scale_b > 0:
                tmp_d = tempfile.mktemp(suffix=".docx")
                build_compact_docx(paragraphs, rng_b, cw_cm, ch_cm, page_margin_cm, tmp_d)
                modify_docx_font_sizes(tmp_d, manual_scale_b * (scale_b / 100), tmp_d)
                pdf_b = tempfile.mktemp(suffix=".pdf")
                docx_to_pdf(tmp_d, pdf_b)
                os.unlink(tmp_d)
            else:
                auto_sb, pdf_b = estimate_font_scale_and_render(
                    docx_path, paragraphs, rng_b,
                    cw_cm, ch_cm, page_margin_cm,
                    target_pages=1, max_target_pages=2,
                )
                if scale_b != 100:
                    adj = auto_sb * (scale_b / 100.0)
                    adj = max(0.35, min(5.0, adj))
                    if abs(adj - auto_sb) > 0.03:
                        os.unlink(pdf_b)
                        tmp_d = tempfile.mktemp(suffix=".docx")
                        build_compact_docx(paragraphs, rng_b, cw_cm, ch_cm, page_margin_cm, tmp_d)
                        modify_docx_font_sizes(tmp_d, adj, tmp_d)
                        pdf_b = tempfile.mktemp(suffix=".pdf")
                        docx_to_pdf(tmp_d, pdf_b)
                        os.unlink(tmp_d)

            section_pdfs = [pdf_a, pdf_b]
            labels = ["Section A", "Section B"]
        else:
            rng_full = [(0, len(paragraphs))]

            if manual_scale_a > 0:
                tmp_d = tempfile.mktemp(suffix=".docx")
                build_compact_docx(paragraphs, rng_full, cw_cm, ch_cm, page_margin_cm, tmp_d)
                modify_docx_font_sizes(tmp_d, manual_scale_a * (scale_a / 100), tmp_d)
                pdf_a = tempfile.mktemp(suffix=".pdf")
                docx_to_pdf(tmp_d, pdf_a)
                os.unlink(tmp_d)
            else:
                auto_sa, pdf_a = estimate_font_scale_and_render(
                    docx_path, paragraphs, rng_full,
                    cw_cm, ch_cm, page_margin_cm,
                    target_pages=1, max_target_pages=2, min_scale=0.35,
                )
                if scale_a != 100:
                    adj = auto_sa * (scale_a / 100.0)
                    adj = max(0.35, min(5.0, adj))
                    if abs(adj - auto_sa) > 0.03:
                        os.unlink(pdf_a)
                        tmp_d = tempfile.mktemp(suffix=".docx")
                        build_compact_docx(paragraphs, rng_full, cw_cm, ch_cm, page_margin_cm, tmp_d)
                        modify_docx_font_sizes(tmp_d, adj, tmp_d)
                        pdf_a = tempfile.mktemp(suffix=".pdf")
                        docx_to_pdf(tmp_d, pdf_a)
                        os.unlink(tmp_d)

            section_pdfs = [pdf_a]
            labels = ["Exam"]

        output_pdf = tempfile.mktemp(suffix=".pdf")
        impose_grid_pages(
            section_pdfs, output_pdf,
            cols=cols, rows=rows,
            margin_mm=margin_mm, gap_mm=gap_mm,
            cut_marks=True,
            fill_mode=False,
            labels=labels,
        )
        return output_pdf

    finally:
        for f in [pdf_a, pdf_b]:
            if f and os.path.exists(f):
                try:
                    os.unlink(f)
                except Exception:
                    pass


def image_to_pdf(image_path: str, output_pdf: str) -> str:
    """Convert a single image to a PDF page using Pillow."""
    img = Image.open(image_path)
    if img.mode != "RGB":
        img = img.convert("RGB")
    img.save(output_pdf, "PDF", resolution=300)
    return output_pdf


def process_single_image(image_path: str, output_pdf: str) -> str:
    """Process a single exam image: convert to PDF with proper sizing on A4."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    if not image_path or not os.path.exists(image_path):
        raise FileNotFoundError(f"Image not found: {image_path}")

    img = Image.open(image_path)
    img_w, img_h = img.size

    pw, ph = A4
    scale = min(pw / img_w, ph / img_h) * 0.9
    dw = img_w * scale
    dh = img_h * scale

    c = canvas.Canvas(output_pdf, pagesize=(pw, ph))
    x = (pw - dw) / 2
    y = (ph - dh) / 2
    c.drawImage(image_path, x, y, width=dw, height=dh, preserveAspectRatio=True)
    c.save()
    return output_pdf


def generate_pdf_previews(pdf_path: str, max_w: int = 900, dpi: int = 72) -> list[str]:
    """Convert PDF pages to base64 preview PNGs."""
    tmpdir = tempfile.mkdtemp(prefix="preview_")
    try:
        prefix = os.path.join(tmpdir, "p")
        subprocess.run(
            ["pdftoppm", "-png", f"-r", str(dpi), pdf_path, prefix],
            capture_output=True, timeout=30,
        )
        imgs = sorted(glob.glob(f"{prefix}*.png"))
        previews = []
        for p in imgs:
            img = Image.open(p)
            w, h = img.size
            scale = max_w / w
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, "PNG")
            previews.append(base64.b64encode(buf.getvalue()).decode())
        return previews
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════════
# Text-based PDF rendering (replaces LibreOffice/DOCX pipeline)
# Uses Reportlab platypus — no system dependencies
# ═══════════════════════════════════════════════════════════════════

def _ocr_text_to_html(text: str) -> str:
    """Convert OCR text (with **bold** markers) to Reportlab-compatible HTML."""
    escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    return escaped.replace("\n", "<br/>\n")


def render_text_to_pdf(
    text: str,
    output_pdf: str,
    cell_width_mm: float,
    cell_height_mm: float,
    font_size_pt: float = 10,
    margin_mm: float = 3,
) -> int:
    """
    Render OCR text directly to a PDF page sized to one grid cell.
    Returns the number of pages in the generated PDF.
    Uses Reportlab platypus — no LibreOffice, no HTML, no subprocess.
    """
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.lib.pagesizes import mm as rl_mm
    import fitz

    page_w = cell_width_mm * rl_mm
    page_h = cell_height_mm * rl_mm

    styles = {
        "center-bold": ParagraphStyle(
            "center-bold", fontName="Helvetica-Bold",
            fontSize=font_size_pt, leading=font_size_pt * 1.3,
            alignment=TA_CENTER, spaceBefore=0, spaceAfter=2,
        ),
        "normal": ParagraphStyle(
            "normal", fontName="Helvetica",
            fontSize=font_size_pt, leading=font_size_pt * 1.3,
            alignment=TA_LEFT, spaceBefore=0, spaceAfter=1,
        ),
        "section-header": ParagraphStyle(
            "section-header", fontName="Helvetica-Bold",
            fontSize=font_size_pt + 1, leading=(font_size_pt + 1) * 1.3,
            alignment=TA_CENTER, spaceBefore=4, spaceAfter=2,
        ),
        "instruction": ParagraphStyle(
            "instruction", fontName="Helvetica-Bold",
            fontSize=font_size_pt, leading=font_size_pt * 1.3,
            alignment=TA_LEFT, spaceBefore=2, spaceAfter=2,
        ),
    }

    html = _ocr_text_to_html(text)
    lines = html.split("\n")
    flowables = []
    line_idx = 0

    for raw in lines:
        stripped = raw.strip()
        if not stripped:
            flowables.append(Spacer(1, font_size_pt * 0.5))
            line_idx += 1
            continue

        is_section = re.match(r"<b>section\s", stripped, re.I)
        is_header = line_idx < 3 and not is_section and not stripped.startswith("SUBJECT:")
        is_instruction = stripped.startswith("<b>INSTRUCTION")

        if is_header:
            style = styles["center-bold"]
        elif is_section:
            style = styles["section-header"]
        elif is_instruction:
            style = styles["instruction"]
        elif stripped.startswith("SUBJECT:"):
            style = styles["normal"]
        else:
            style = styles["normal"]

        flowables.append(Paragraph(stripped, style))
        line_idx += 1

    margin = margin_mm * rl_mm
    doc = SimpleDocTemplate(
        output_pdf,
        pagesize=(page_w, page_h),
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin,
    )
    doc.build(flowables)

    doc_fitz = fitz.open(output_pdf)
    pages = doc_fitz.page_count
    doc_fitz.close()
    return pages


def generate_section_pdf_binary(
    text: str,
    cell_width_mm: float,
    cell_height_mm: float,
    target_pages: int = 1,
    min_font: float = 5,
    max_font: float = 14,
    tolerance: float = 0.25,
) -> str:
    """
    Binary search for largest font size that fits text in ≤ target_pages
    at cell size. Returns path to generated PDF.

    0 LibreOffice calls — pure Reportlab + PyMuPDF.
    """
    low, high = min_font, max_font
    best_pdf = None
    best_font = max_font

    while high - low > tolerance:
        mid = (low + high) / 2
        tmp = tempfile.mktemp(suffix=".pdf")
        pages = render_text_to_pdf(text, tmp, cell_width_mm, cell_height_mm, mid)

        if pages <= target_pages:
            if best_pdf:
                os.unlink(best_pdf)
            best_pdf = tmp
            best_font = mid
            low = mid
        else:
            os.unlink(tmp)
            high = mid

    if best_pdf is None:
        best_pdf = tempfile.mktemp(suffix=".pdf")
        render_text_to_pdf(text, best_pdf, cell_width_mm, cell_height_mm, min_font)

    return best_pdf


def _split_ocr_into_sections(text: str) -> dict:
    """Split OCR text into Section A and Section B by finding **SECTION B** marker."""
    m = re.split(r"(?i)\*\*section\s+b", text, maxsplit=1)
    if len(m) > 1:
        before = m[0].strip()
        after = ("**SECTION B" + m[1]).strip()
        return {"Section A": before, "Section B": after}
    return {"Exam": text.strip()}


def impose_from_text(
    text_or_sections: str | dict,
    output_pdf_path: str,
    cols: int = 3,
    rows: int = 2,
    margin_mm: float = 2.5,
    gap_mm: float = 2,
    target_pages_per_section: int = 1,
    labels: list | None = None,
):
    """
    Create a 3×2 imposed PDF directly from OCR text.
    No LibreOffice or DOCX involved.

    text_or_sections: raw OCR string (auto-split on Section B marker)
                      or pre-split dict {"Section A": "...", "Section B": "..."}
    """
    from reportlab.lib.units import mm as rl_mm

    if isinstance(text_or_sections, str):
        sections_text = _split_ocr_into_sections(text_or_sections)
    else:
        sections_text = text_or_sections

    pw, ph = rl_landscape(A4)
    mg = margin_mm * rl_mm
    gp = gap_mm * rl_mm
    cell_w = (pw - 2 * mg - (cols - 1) * gp) / cols
    cell_h = (ph - 2 * mg - (rows - 1) * gp) / rows
    cell_w_mm = cell_w / rl_mm
    cell_h_mm = cell_h / rl_mm

    section_pdfs = []
    try:
        for name, text in sections_text.items():
            pdf = generate_section_pdf_binary(
                text, cell_w_mm, cell_h_mm,
                target_pages=target_pages_per_section,
            )
            section_pdfs.append(pdf)

        impose_grid_pages(
            section_pdfs, output_pdf_path,
            cols=cols, rows=rows,
            margin_mm=margin_mm, gap_mm=gap_mm,
            cut_marks=True, fill_mode=False,
            labels=labels or list(sections_text.keys()),
        )
        return output_pdf_path
    finally:
        for p in section_pdfs:
            if p and os.path.exists(p):
                try:
                    os.unlink(p)
                except Exception:
                    pass
